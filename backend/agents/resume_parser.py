from .base import BaseAgent
from .prompts import RESUME_PARSER_PROMPT
from models.schemas import ParsedResume
import pdfplumber
import fitz  # pymupdf
from docx import Document
from io import BytesIO

class ResumeParserAgent(BaseAgent):
    @property
    def model(self) -> str:
        return self.settings.extraction_model
    
    @property
    def temperature(self) -> float:
        return self.settings.parsing_temp
    
    async def execute(self, file_content: bytes, filename: str) -> ParsedResume:
        text = self._extract_text(file_content, filename)
        prompt = RESUME_PARSER_PROMPT.format(resume_text=text)
        
        response = self._call_llm(
            prompt,
            system_prompt="You are a precise resume parser. Return only valid JSON."
        )
        
        data = self._parse_json(response)
        return ParsedResume(**data)
    
    def _extract_text(self, content: bytes, filename: str) -> str:
        ext = filename.lower().split('.')[-1]
        
        if ext == 'pdf':
            return self._extract_pdf(content)
        elif ext in ['docx', 'doc']:
            return self._extract_docx(content)
        elif ext == 'txt':
            return content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_pdf(self, content: bytes) -> str:
        text_parts = []
        
        # Try pdfplumber first (better for text-based PDFs)
        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception:
            pass
        
        # Fallback to pymupdf if pdfplumber fails
        if not text_parts:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
            except Exception as e:
                raise ValueError(f"Could not extract text from PDF: {e}")
        
        return "\n".join(text_parts)
    
    def _extract_docx(self, content: bytes) -> str:
        doc = Document(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return "\n".join(paragraphs)
